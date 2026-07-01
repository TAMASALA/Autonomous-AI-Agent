from docx import Document
import uuid

def create_doc(content):

    doc=Document()

    for line in content.split("\n"):

        if line.startswith("# "):
            doc.add_heading(line.replace("# ",""),1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ",""),2)

        else:
            doc.add_paragraph(line)

    filename=f"generated_docs/{uuid.uuid4()}.docx"

    doc.save(filename)

    return filename